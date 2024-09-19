import os
import json
import logging
import time
import re
import whisper
import base64
import psycopg2
from PIL import Image, ImageChops
import urllib.parse
import vertexai.preview.generative_models as generative_models
import vertexai
from vertexai.preview.generative_models import GenerativeModel
from docx import Document
import urllib.parse
from docx.shared import Pt
from docx.shared import Inches
from flask import jsonify, request, Flask
from datetime import datetime
import os
import cv2
from google.cloud import vision
from google.oauth2 import service_account
from google.cloud.vision_v1 import ImageAnnotatorClient
from transformers import pipeline
 
app = Flask(__name__)
# Initialize the summarization pipeline
summarizer = pipeline("summarization", model="t5-small")
 
target_path = "./Logs"
credentials = service_account.Credentials.from_service_account_file('config.json')
client = vision.ImageAnnotatorClient(credentials=credentials)
 
# Initialize logger
logs_dir = os.path.join(target_path, 'transcribeEase_logs')
if not os.path.exists(logs_dir):
    os.makedirs(logs_dir)
    print(f"transcribeEase_logs directory created successfully at '{logs_dir}'.")
   
application_name = "transcribeEase_logs"
log_file_path = os.path.join(logs_dir, f"{datetime.now().strftime('%Y-%m-%d')}_{application_name}")
if not os.path.exists(log_file_path):
    os.makedirs(log_file_path)
    print(f"transcribeEase_logs directory created successfully at '{log_file_path}'.")
 
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
 
# Create handlers for each log level
file_handlers = {}
levels = ['error', 'info', 'debug', 'critical']
for level in levels:
    file_handler = logging.FileHandler(os.path.join(log_file_path, f"{level}.log"))
    file_handler.setLevel(logging.getLevelName(level.upper()))
    file_handler.setFormatter(formatter)
    file_handlers[level] = file_handler
 
    logger.addHandler(file_handler)
 
global transcript_content
with open('postgrace_credential.json', 'r') as config_file:
    config = json.load(config_file)
    dbname = config.get('dbname')
    user = config.get('user')
    password = config.get('password')
    host = config.get('host')
    port = config.get('port')
 
 
# transcribeEase_blueprint = Blueprint('transcribeEase', __name__)
 
with open('config.json', 'r') as config_file:
    config = json.load(config_file)
    project_id = config.get('project_id')
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "config.json"
project_id = project_id
print(project_id)
 
output_folder = "./captured_screenshots"
screenshot_folder = os.path.join(output_folder, "screenshots")
ocr_result_folder = os.path.join(output_folder, "ocr_results")
 
 
def are_images_similar(img1, img2, threshold=5):
    """Compare two images to check if they are similar."""
    diff = ImageChops.difference(img1, img2)
    if diff.getbbox() is None:
        return True  # Images are identical
    else:
        return False  # Images are different
   
def perform_ocr(image_content):
    """
    Perform OCR on the given image content and extract text content using Google Cloud Vision API.
 
    Args:
    - image_content (bytes): The binary content of the image.
    - context (str): Additional context or information related to the image.
 
    Returns:
    - str: The extracted text content along with the provided context.
    """
    image = vision.Image(content=image_content)
    response = client.text_detection(image=image)
    texts = response.text_annotations
 
    if texts:
        return texts[0].description
    else:
        return "No text found in the image"
   
def capture_screenshots_and_perform_ocr(video_file, output_folder, interval=10, max_screenshots=10):
    """
    Capture screenshots from a video file at regular intervals, perform OCR on each screenshot, and save the OCR results.
 
    Args:
    - video_file (str): Path to the video file.
    - output_folder (str): Path to the folder where screenshots and OCR results will be saved.
    - interval (int): Interval in milliseconds between consecutive screenshots.
    - max_screenshots (int): Maximum number of screenshots to capture.
 
    Returns:
    - List of tuples: List of tuples containing (screenshot_path, ocr_result_path).
    """
    try:
        os.makedirs(output_folder, exist_ok=True)
        cap = cv2.VideoCapture(video_file)
        frame_count = 0
        prev_frame = None
        screenshot_count = 0
        results = []
 
        while cap.isOpened() and screenshot_count < max_screenshots:
            success, frame = cap.read()
 
            if success:
                frame_count += 1
 
                if frame_count * interval >= 1000:
                    current_frame = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
 
                    if prev_frame is not None and are_images_similar(prev_frame, current_frame):
                        print("Duplicate frame. Skipping...")
                    else:
                        screenshot_path = os.path.join(output_folder, f"screenshot_{screenshot_count}.jpg")
                        current_frame.save(screenshot_path)
                        print(f"Screenshot saved: {screenshot_path}")
 
                        with open(screenshot_path, 'rb') as image_file:
                            image_content = image_file.read()
                        text_content = perform_ocr(image_content)
 
                        ocr_result_path = os.path.join(output_folder, f"ocr_result_{screenshot_count}.txt")
                        with open(ocr_result_path, 'w') as f:
                            f.write(text_content)
 
                        results.append((screenshot_path, ocr_result_path))
                        screenshot_count += 1
 
                    prev_frame = current_frame
                    frame_count = 0
 
        cap.release()
        cv2.destroyAllWindows()
        print("Screenshot capture and OCR completed.")
 
        return results
 
    except Exception as e:
        print(f"An error occurred: {e}")
        return []
   
 
 
@app.route('/getEntireTableVerizontranscribe', methods=['GET'])
def get_entire_tabletranscribe():
    try:
        conn = psycopg2.connect(
            dbname=dbname,
            user=user,
            password=password,
            host=host,
            port = port
        )
        cursor = conn.cursor()
        table_name = "audio_transcript"
        cursor.execute("SELECT * FROM audio_transcript")
        rows = cursor.fetchall()
        cursor.execute("SELECT COUNT(*) FROM audio_transcript")
        count = cursor.fetchone()[0]
        cursor.execute("SELECT column_name FROM information_schema.columns WHERE table_name = %s", (table_name,))
        columns = cursor.fetchall()
        column_names = [col[0] for col in columns]
       
        print(count)
        conn.close()
        data = []
        for row in rows:
            row_dict = {}
            for idx, column_name in enumerate(column_names):
                if column_name == 'audio':
                    if row[idx] is not None:
                        row_dict[column_name] = base64.b64encode(row[idx]).decode('utf-8')
                else:
                    row_dict[column_name] = row[idx]
            data.append(row_dict)
        return jsonify({"data": data})
    except psycopg2.Error as e:
        return jsonify({'error': 'An error occurred while fetching data from the database'}), 500
   
 
@app.route('/upload_audio_file', methods=['POST'])
def upload_audio_file():
   
    doc = Document()
   
    try:
        # audio_file = request.files.get('file')
        video_file = request.files.get('file')
       
        # input_filename1 = audio_file.filename
        # file_split1 = input_filename1.split(".")
        input_filename = video_file.filename
        file_split = input_filename.split(".")
        # output_filename1 = file_split1[0]
        output_filename = file_split[0]
        current_date = datetime.now().date()
        temp_file_path = "/tmp/temp_audio_files"
        video_file.save(temp_file_path)
        start_time = time.time()
        start_transcript_time = time.time()
       
 
        # Capture screenshots from the video file and perform OCR
       
 
       
        logging.info("wishper started generating successfully")
        model = whisper.load_model("medium")
        result = model.transcribe(temp_file_path)
        transcript = result["text"]
        end_time = time.time()
        s1 = time.time()
        segments = result['segments']
        e1 = time.time()
 
        print(e1-s1)
       
        # Print out the transcriptions with timestamps
        for segment in segments:
            start = segment['start']
            end = segment['end']
            text = segment['text']
            segment_system_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
           
           
            transcript_entry = f"[{segment_system_time}] Transcript generated in [{start:.2f}s - {end:.2f}s]: {text}"
            print(transcript_entry)
 
            # Add the transcript entry to the document
            doc.add_paragraph(transcript_entry)
        # screenshot_folder = f"{output_folder}_transcript_{current_date}"
        # ocr_result_folder = f"{output_folder}_ocr_result_{current_date}"
       
        transcript = result["text"]
       
        end_transcript_time = time.time()
       
        transcript_generation_time = round(end_transcript_time - start_transcript_time)
        # output_folder=screenshot_folder
        result = capture_screenshots_and_perform_ocr(temp_file_path, output_folder)
        print(result)
       
        os.remove(temp_file_path)
        total_time_audio = round(end_time-start_time)
        logging.info("transcript generated successfully")
    except Exception as e:
        # total_time_audio = None
        logger.error(f"Error occurred in transprit generation: {str(e)}")
   
    # transcript_content = ""
    # transcript = "Hi, thank you for talking to me. I'm a social media agent. Can I have your name, please? Hello? Hi, thank you so much for talking to me. I'm a social media agent. Can I have your name, please? I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. I'm a social media agent. So, there's some tips on the volunteer and see you in the class. What's one guy? You need to turn your data off. Turn your data off? Yes, you need to turn your data off. Because you don't have enough space to put it on. That's why the credits just keep going. Okay. Alright, so go ahead. Go ahead, Ms. Alexander. Hello? Alexander? Yes, I'm saying that's what happened to the credits. No, not yet. Alright, keep it going. Thank you so much. You're welcome."
    try:
        vertexai.init(project=project_id, location="asia-southeast1")
        config = {
                    "max_output_tokens": 8192,
                    "temperature": 0.2,
                    "top_p": 1
            }
        model = GenerativeModel("gemini-1.0-pro-001")
        chat = model.start_chat()
        code = chat.send_message(f"""\"{transcript}"
    Generate a code for Web Sequence Diagram of above transcript.""", generation_config=config, safety_settings={})
        print(code.text)
   
        response = chat.send_message(f"""{transcript}"
   
    Generate a well-structured document and markdown headings for Angular and proper numbering for each heading. This document should encompass the following sections with specific details:
   
    0- Heading: provide a two words heading that conveys the essence of recording
    1- Main Topics : Provide a concise four-line explanation of Identify the overarching themes and the specific subjects discussed during the session.
    2- Subtopics : Provide a concise four-line explanation of Identify the overarching themes and the specific subjects discussed during the session.
    3 - Key Concepts and Definitions: Highlight important terms, concepts, and their definitions to ensure clarity and understanding.
    4- Pain points: Provide a concise four-line pain point of Document discussed during the session.
    5- Detailed Processe: Provide a concise four-line explanation of Document step-by-step procedures, workflows, and processes discussed during the session.
    6- Best Practices: Extract insights on best practices, tips, and lessons learned from experiences shared during the session.
    7- Case Studies and Examples: Note any relevant case studies, examples, or real-life scenarios used to illustrate concepts or principles.
    8- Tools and Technologies: Record mentions of tools, software, technologies, or methodologies discussed and their applications.
    9- Action Items and Next Steps: Record any action items, tasks, or next steps identified during the session to ensure follow-up and implementation.
    10- Key feature : provide the key feature of input transcript.
    Always keep in mind if anything is not available in recording write \"None\".And generate maximum points of each heading as possible.""", generation_config=config, safety_settings={
            generative_models.HarmCategory.HARM_CATEGORY_HATE_SPEECH: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            generative_models.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            generative_models.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
                generative_models.HarmCategory.HARM_CATEGORY_HARASSMENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
                })
        print(response.text)
   
   
        faq = chat.send_message(f"""\"{transcript}"
   
    Generate a maximum number of FAQ questions with answer related to above transcript as you can .I am giving you the format of FAQ.
    Q:
    A:""", generation_config=config, safety_settings={
            generative_models.HarmCategory.HARM_CATEGORY_HATE_SPEECH: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            generative_models.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            generative_models.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            generative_models.HarmCategory.HARM_CATEGORY_HARASSMENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
                })
        print(faq.text)
        summary = chat.send_message(f"""\"{transcript}"
   
    Your task is to write above transcript in proper markdown headings for angular.
    1.Heading : provide a two words heading that conveys the essence of recording.
   
    2.Summary : provide a summary of input trancript.
   
    3.Description : in this section directly write whole input trancript that i will provide you dont do anything else.
   
    4.Conclusion : provide a proper conclusion of the above transcript.""", generation_config=config, safety_settings={
            generative_models.HarmCategory.HARM_CATEGORY_HATE_SPEECH: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            generative_models.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            generative_models.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
            generative_models.HarmCategory.HARM_CATEGORY_HARASSMENT: generative_models.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE,
                })
   
        print(summary.text)
        transcript_content = f"KT Session : {response.text}\n\n{faq.text}\n\n Detailed Transcript : \n{summary.text}"
 
        # transcript_content = f"KT Session : {response.text}\n\n{faq.text}\n\n Detailed Transcript : \n{summary.text}"
    except Exception as e:
        logger.error(f"Error occurred in GCp models: {str(e)}")
       
    def getSequenceDiagram(text, outputFile, style='default'):
                request = {}
                request["message"] = text
                request["style"] = style
                request["apiVersion"] = "1"
 
                url = urllib.parse.urlencode(request).encode('utf-8')
 
                with urllib.request.urlopen("https://www.websequencediagrams.com/", url) as f:
                    line = f.readline().decode('utf-8')
 
                expr = re.compile("(\?(img|pdf|png|svg)=[a-zA-Z0-9]+)")
                m = expr.search(line)
 
                if m is None:
                    return False
 
                urllib.request.urlretrieve("https://www.websequencediagrams.com/" + m.group(0), outputFile)
                return True
    href = None
    try:
        flowcode = code.text
        style = "qsd"
        text = "alice->bob: authentication request\nbob-->alice: response"
        pngFile = "diagram.png"
        getSequenceDiagram(flowcode, pngFile, style)
    except Exception as e:
        logger.error(f"{str(e)}")
       
    # output_folder = f"output_doc/{output_filename}_transcript_{current_date}.docx"
   
 
    # document_output_folder = f"{output_filename}_transcript_{current_date}.docx"
 
    # Create the TXT document
   
    # Create the document
   
       
        # Add content to the document
    doc.add_paragraph(transcript_content) # Add the transcript
    doc.add_picture(pngFile, width=Inches(4))  
   
       
        # Loop through each screenshot and its OCR result
    max_screenshots = 5  # Example value, change as needed
    for i in range(max_screenshots):
           
           
        screenshot_path = os.path.join(output_folder, f'screenshot_{i}.jpg')
        ocr_result_path = os.path.join(output_folder, f'ocr_result_{i}.txt')
        if not os.path.exists(screenshot_path) or not os.path.exists(ocr_result_path):
            break
 
        screenshot = Image.open(screenshot_path)
        ocr_content = ""
        with open(ocr_result_path, 'r') as f:
            ocr_content = f.read()
        screenshot_capture_time = datetime.now().strftime("%H:%M:%S")
        doc.add_heading(f'Screenshot {i+1}', level=2)
        doc.add_paragraph(f'Screenshot captured at: {screenshot_capture_time}')
        doc.add_picture(screenshot_path, width=Inches(4))
        doc.add_paragraph(ocr_content)
       
       
        # Add the sequence diagram
    # Add screenshots, OCR content, and explanations to the document
        # Generate a summary of the OCR content
        summary = summarizer(ocr_content, max_length=450, min_length=150, do_sample=False)[0]['summary_text']
       
        # Add the summary to the document
        doc.add_heading('Summary', level=3)
        doc.add_paragraph(summary)
       
 
            # Save the document
    document_output_folder = f"{output_filename}_transcript_{current_date}.docx"
    doc.save(document_output_folder)
       
        # Convert the saved document to base64
    with open(document_output_folder, "rb") as f:
        doc_data = f.read()
       
    b64 = base64.b64encode(doc_data).decode()
       
        # Construct the data URI for downloading the document
    href = f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}"
       
    print("File downloaded successfully.")
    current_time = datetime.now().strftime("%H:%M")
 
    try:
        conn = psycopg2.connect(
            dbname=dbname,
            user=user,
            password=password,
            host=host,
            port = port
        )
        cursor = conn.cursor()
        cursor.execute('''CREATE TABLE IF NOT EXISTS audio_transcript (
                            "fileName" TEXT PRIMARY KEY,
                            transcript TEXT,
                            month TEXT,
                            atdate TEXT,
                            attime TEXT,
                            downloadlink TEXT,
                            total_time_audio INTEGER
                        )''')
 
        conn.commit()
 
        audio_binary = video_file.read()
        print(audio_binary)
        atTime = datetime.now()
        month = datetime.now().strftime("%B")
        atDate = datetime.now().date()
        cursor.execute('''INSERT INTO audio_transcript
                  ("fileName", transcript, month, atdate, attime, total_time_audio, downloadlink)
                  VALUES (%s, %s, %s, %s, %s, %s, %s)''',
                  (output_filename, transcript, month, atDate, atTime, total_time_audio, href))
        conn.commit()
    except Exception as e:
        print("An error occurred:", e)          
    return jsonify({"transcript":transcript,'doc': href,"name":f"{output_filename}_transcript_{current_date}.docx","date":current_date,"time":current_time})
 
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=7800)